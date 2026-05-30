"""
Perceiver Agent - Multi-Modal Sensory Input Processing.

The Perceiver agent provides:
- Multi-modal input handling (text, image, audio, video)
- Feature extraction and preprocessing
- Sensory data normalization and encoding
- Cross-modal correlation and fusion
- Input quality assessment and filtering

Named for the ability to perceive and process sensory information from multiple sources.
"""

import asyncio
import base64
import hashlib
import tempfile
from datetime import UTC, datetime
from typing import Any, ClassVar

import structlog
from swarms import Agent

from heretek_swarm.actors.base import ActorMessage, AgentActor

# Session 44: Collective Learning Integration
from heretek_swarm.actors.mixins import (
    DeliberationMixin,
    HealthReportingMixin,
    LearningMixin,
    MemoryMixin,
    PatternMixin,
    ValidationMixin,
)
from heretek_swarm.actors.perceiver.types import ModalityType
from heretek_swarm.actors.validation import validate_message
from heretek_swarm.collective.learning import PatternExtractor

# Session 44: Consensus Integration
from heretek_swarm.consensus.swarm_deliberation import SwarmDeliberationEngine

# Session 44: Memory Optimization Integration
from heretek_swarm.memory.access_patterns import AccessPatternAnalyzer

# Session 44: Zero-Trust Validation
from heretek_swarm.security.zero_trust import ZeroTrustValidator
from heretek_swarm.workflow.validator import ValidationError

logger = structlog.get_logger("PerceiverAgent")

# Module-level constant for repeated MIME type
_OCTET_STREAM_MIME = "application/octet-stream"
_DATA_URL_PREFIX = "data:"


class PerceiverAgent(
    HealthReportingMixin,
    ValidationMixin,
    DeliberationMixin,
    PatternMixin,
    MemoryMixin,
    LearningMixin,
    AgentActor,
):
    """
    Perceiver Agent - Multi-Modal Sensory Input Processing Specialist.

    The Perceiver is responsible for:
    - Receiving and classifying multi-modal inputs
    - Extracting features from sensory data
    - Normalizing and encoding inputs for downstream agents
    - Correlating information across modalities
    - Assessing input quality and filtering noise

    Sensory Processing Workflow:
    1. Receive input with modality classification
    2. Validate input format and quality
    3. Extract modality-specific features
    4. Normalize and encode for swarm consumption
    5. Store processed data in Historian memory
    6. Notify relevant agents of new input
    """

    def __init__(
        self,
        agent_id: str = "perceiver",
        name: str = "Perceiver",
        description: str = "Multi-modal sensory input processing specialist",
        swarms_agent: Agent | None = None,
        max_input_size_mb: int = 50,
        feature_cache_size: int = 1000,
        enable_cross_modal: bool = True,
        pattern_extractor: PatternExtractor | None = None,
        deliberation_engine: SwarmDeliberationEngine | None = None,
        access_analyzer: AccessPatternAnalyzer | None = None,
        zero_trust_validator: ZeroTrustValidator | None = None,
        **kwargs,
    ) -> None:
        """
        Initialize the Perceiver agent.

        Args:
            agent_id: Unique identifier
            name: Human-readable name
            description: Agent description
            swarms_agent: Optional Swarms Agent for LLM capabilities
            max_input_size_mb: Maximum input size in megabytes
            feature_cache_size: Maximum cached feature entries
            enable_cross_modal: Enable cross-modal correlation
            **kwargs: Additional arguments
        """
        super().__init__(
            agent_id=agent_id,
            name=name,
            description=description,
            topics=[
                "sensory-input",
                "multi-modal",
                "feature-extraction",
                "preprocessing",
            ],
            capabilities=[
                "text-processing",
                "image-analysis",
                "audio-processing",
                "video-analysis",
                "document-parsing",
                "sensor-data",
                "feature-extraction",
                "cross-modal-fusion",
            ],
            swarms_agent=swarms_agent,
            **kwargs,
        )

        # Perceiver-specific configuration
        self.max_input_size_mb = max_input_size_mb
        self.feature_cache_size = feature_cache_size
        self.enable_cross_modal = enable_cross_modal

        # Processing statistics
        self.inputs_processed: dict[str, int] = {modality.value: 0 for modality in ModalityType}
        self.total_features_extracted = 0
        self.quality_rejections = 0

        # Feature cache for cross-modal correlation
        self.feature_cache: dict[str, dict[str, Any]] = {}
        self.cross_modal_correlations: list[dict[str, Any]] = []

        # Supported formats per modality
        self.supported_formats: dict[str, list[str]] = {
            ModalityType.TEXT.value: ["txt", "md", "json", "xml", "html"],
            ModalityType.IMAGE.value: ["jpg", "jpeg", "png", "gif", "bmp", "webp", "svg"],
            ModalityType.AUDIO.value: ["mp3", "wav", "ogg", "flac", "aac"],
            ModalityType.VIDEO.value: ["mp4", "avi", "mov", "webm", "mkv"],
            ModalityType.DOCUMENT.value: ["pdf", "doc", "docx", "ppt", "pptx", "xls", "xlsx"],
            ModalityType.SENSOR.value: ["json", "csv", "binary"],
        }

        # Session 44: Collective Learning Integration
        self.pattern_extractor = pattern_extractor or PatternExtractor(
            min_support=3, min_confidence=0.6
        )

        # Session 44: Consensus Integration
        self.deliberation_engine = deliberation_engine or SwarmDeliberationEngine(
            max_rounds=5, consensus_threshold=0.75, min_participants=2
        )

        # Session 44: Memory Optimization Integration
        self.access_analyzer = access_analyzer or AccessPatternAnalyzer()

        # Session 44: Zero-Trust Validation
        self.zero_trust_validator = zero_trust_validator or ZeroTrustValidator()

        # Session 44: Integration state
        self._active_deliberations: dict[str, str] = {}
        self._pattern_emitted: set[str] = set()

        logger.info(f"[{self.agent_id}] Perceiver agent initialized")

    async def initialize(self) -> None:
        """Initialize the Perceiver agent."""
        # Register message handlers with Zero-Trust validation
        self.register_handler("process_input", self._handle_process_input)
        self.register_handler("extract_features", self._handle_extract_features)
        self.register_handler("classify_modality", self._handle_classify_modality)
        self.register_handler("assess_quality", self._handle_assess_quality)
        self.register_handler("get_processing_stats", self._handle_get_processing_stats)
        self.register_handler("correlate_modalities", self._handle_correlate_modalities)

        logger.info(f"[{self.agent_id}] Perceiver initialization complete")

    async def process_message(self, message: ActorMessage) -> None:
        """
        Process incoming messages with exception handling.

        Args:
            message: Actor message to process
        """
        handler = self._message_handlers.get(message.message_type)
        if handler:
            try:
                await handler(message)
            except Exception as e:
                logger.exception(
                    f"[{self.agent_id}] Error processing message {message.message_type}: {e}",  # noqa: G004

                )
                self.error_count += 1
                if message.content.get("reply_to"):
                    await self.send(
                        topic=message.content["reply_to"],
                        content={
                            "message_type": "error_response",
                            "error": str(e),
                            "original_message_type": message.message_type,
                        },
                        correlation_id=message.correlation_id,
                    )
        else:
            logger.warning(f"[{self.agent_id}] No handler for message type: {message.message_type}")

    def _validate_message_content(self, message_type: str, content: dict[str, Any]) -> Any:
        """Validate message content using Pydantic models."""
        try:
            return validate_message(message_type, content)
        except ValidationError as e:
            logger.warning(
                f"[{self.agent_id}] Message validation failed for {message_type}: {e}",  # noqa: G004
                extra={"validation_errors": e.errors()},
            )
            raise ValueError(f"Invalid message format: {e.errors()}") from e
        except KeyError:
            logger.debug(f"[{self.agent_id}] No validator for message type: {message_type}")
            return None

    async def _handle_process_input(self, message: ActorMessage) -> None:
        """
        Process multi-modal input.

        Args:
            message: ActorMessage with content containing:
                - input_data: The raw input (text, base64 data, or URL)
                - modality: Optional modality type (auto-detected if not provided)
                - format: Optional format hint (e.g., 'jpg', 'txt')
                - metadata: Optional metadata dict
                - priority: Optional priority level (1-10)

        Response:
            - input_id: Unique identifier for this input
            - modality: Detected/provided modality
            - features: Extracted features
            - quality_score: Input quality assessment (0-1)
            - timestamp: Processing timestamp
        """
        try:
            content = message.content
            input_data = content.get("input_data")
            modality = content.get("modality")
            format_hint = content.get("format")
            metadata = content.get("metadata", {})
            content.get("priority", 5)

            if input_data is None:
                await self._send_error_response(message, "Missing input_data")
                return

            # Validate input size
            if not self._validate_input_size(input_data):
                self.quality_rejections += 1
                await self._send_error_response(
                    message, f"Input exceeds maximum size ({self.max_input_size_mb}MB)"
                )
                return

            # Auto-detect modality if not provided
            if not modality:
                modality = self._detect_modality(input_data, format_hint)

            # Generate unique input ID
            input_id = self._generate_input_id(input_data, modality)

            # Extract features based on modality
            features = await self._extract_modality_features(input_data, modality, format_hint)

            # Assess quality
            quality_score = self._assess_input_quality(input_data, modality, features)

            # Cache features for cross-modal correlation
            self._cache_features(input_id, modality, features, metadata)

            # Update statistics
            self.inputs_processed[modality] = self.inputs_processed.get(modality, 0) + 1
            self.total_features_extracted += len(features) if features else 0

            # Store in Historian if available
            await self._store_in_historian(input_id, modality, features, metadata)

            # Send response
            await self.send(
                topic=content.get("reply_to", "actor:*"),
                content={
                    "message_type": "input_processed",
                    "input_id": input_id,
                    "modality": modality,
                    "features": features,
                    "quality_score": quality_score,
                    "timestamp": datetime.now(UTC).isoformat(),
                },
                correlation_id=message.correlation_id,
            )

            logger.info(
                f"[{self.agent_id}] Input processed: {input_id[:8]}...",  # noqa: G004
                extra={"modality": modality, "quality": quality_score},
            )

        except Exception as e:
            logger.exception(f"[{self.agent_id}] Input processing failed: {e}")
            await self._send_error_response(message, f"Input processing failed: {e}")

    def _validate_input_size(self, input_data: Any) -> bool:
        """Validate input does not exceed maximum size."""
        try:
            max_bytes = self.max_input_size_mb * 1024 * 1024
            if isinstance(input_data, str):
                return len(input_data.encode()) <= max_bytes
            if isinstance(input_data, bytes):
                return len(input_data) <= max_bytes
            if isinstance(input_data, dict):
                import json

                return len(json.dumps(input_data).encode()) <= max_bytes
            return True  # Assume valid for other types
        except Exception as e:
            logger.debug("perceiver_validation_failed", error=str(e))
            return True  # Fail open on validation errors

    # Modality detection lookup tables
    _IMAGE_FORMATS: ClassVar[frozenset[str]] = frozenset(
        ["jpg", "jpeg", "png", "gif", "bmp", "webp", "svg"]
    )
    _AUDIO_FORMATS: ClassVar[frozenset[str]] = frozenset(
        ["mp3", "wav", "ogg", "flac", "aac"]
    )
    _VIDEO_FORMATS: ClassVar[frozenset[str]] = frozenset(
        ["mp4", "avi", "mov", "webm", "mkv"]
    )
    _DOCUMENT_FORMATS: ClassVar[frozenset[str]] = frozenset(
        ["pdf", "doc", "docx", "ppt", "pptx", "xls", "xlsx"]
    )

    def _detect_modality(self, input_data: Any, format_hint: str | None = None) -> str:
        """Auto-detect input modality."""
        if format_hint:
            modality = self._detect_from_format_hint(format_hint)
            if modality:
                return modality
        return self._detect_from_content(input_data)

    def _detect_from_format_hint(self, format_hint: str) -> str | None:
        """Detect modality from format hint string."""
        fmt = format_hint.lower()
        if fmt in self._TEXT_FORMATS:
            return ModalityType.TEXT.value
        if fmt in self._IMAGE_FORMATS:
            return ModalityType.IMAGE.value
        if fmt in self._AUDIO_FORMATS:
            return ModalityType.AUDIO.value
        if fmt in self._VIDEO_FORMATS:
            return ModalityType.VIDEO.value
        if fmt in self._DOCUMENT_FORMATS:
            return ModalityType.DOCUMENT.value
        return None

    def _detect_from_content(self, input_data: Any) -> str:
        """Detect modality from content inspection."""
        if isinstance(input_data, str):
            return self._detect_string_modality(input_data)
        if isinstance(input_data, bytes):
            return self._detect_bytes_modality(input_data)
        if isinstance(input_data, dict):
            return ModalityType.SENSOR.value
        return ModalityType.TEXT.value

    def _detect_string_modality(self, data: str) -> str:
        """Detect modality from string content."""
        if data.startswith(_DATA_URL_PREFIX):
            mime_type = data.split(":")[1].split(";")[0]
            if "image" in mime_type:
                return ModalityType.IMAGE.value
            if "audio" in mime_type:
                return ModalityType.AUDIO.value
            if "video" in mime_type:
                return ModalityType.VIDEO.value
        return ModalityType.TEXT.value

    @staticmethod
    def _detect_bytes_modality(data: bytes) -> str:
        """Detect modality from bytes magic numbers."""
        if data.startswith(b"\xff\xd8\xff"):
            return ModalityType.IMAGE.value  # JPEG
        if data.startswith(b"\x89PNG"):
            return ModalityType.IMAGE.value  # PNG
        if data.startswith(b"RIFF") and data[8:12] == b"WAVE":
            return ModalityType.AUDIO.value  # WAV
        return ModalityType.TEXT.value

    def _generate_input_id(self, input_data: Any, modality: str) -> str:
        """Generate unique input identifier."""
        # Create hash of input data for deduplication
        if isinstance(input_data, str):
            data_bytes = input_data.encode()
        elif isinstance(input_data, bytes):
            data_bytes = input_data
        else:
            import json

            data_bytes = json.dumps(input_data, sort_keys=True).encode()

        hash_digest = hashlib.sha256(data_bytes).hexdigest()[:16]
        timestamp = datetime.now(UTC).strftime("%Y%m%d%H%M%S")
        return f"input_{modality}_{timestamp}_{hash_digest}"

    async def _extract_modality_features(
        self, input_data: Any, modality: str, format_hint: str | None
    ) -> dict[str, Any]:
        """Extract features based on modality."""
        try:
            if modality == ModalityType.TEXT.value:
                return self._extract_text_features(input_data)
            if modality == ModalityType.IMAGE.value:
                return await self._extract_image_features(input_data, format_hint)
            if modality == ModalityType.AUDIO.value:
                return await self._extract_audio_features(input_data, format_hint)
            if modality == ModalityType.VIDEO.value:
                return await self._extract_video_features(input_data, format_hint)
            if modality == ModalityType.DOCUMENT.value:
                return await self._extract_document_features(input_data, format_hint)
            if modality == ModalityType.SENSOR.value:
                return self._extract_sensor_features(input_data)
            return {"error": f"Unknown modality: {modality}"}
        except Exception as e:
            logger.error(f"[{self.agent_id}] Feature extraction failed: {e}")
            return {"error": str(e)}

    def _extract_text_features(self, text: str) -> dict[str, Any]:
        """Extract features from text input."""
        if not isinstance(text, str):
            text = str(text)

        # Basic text statistics
        words = text.split()
        sentences = text.replace("!", ".").replace("?", ".").split(".")
        sentences = [s.strip() for s in sentences if s.strip()]

        # Character-level features
        char_count = len(text)
        char_count_no_spaces = len(text.replace(" ", ""))

        # Word-level features
        word_count = len(words)
        avg_word_length = sum(len(w) for w in words) / word_count if word_count > 0 else 0

        # Sentence-level features
        sentence_count = len(sentences)
        avg_sentence_length = word_count / sentence_count if sentence_count > 0 else 0

        # Vocabulary features
        unique_words = {w.lower() for w in words}
        vocabulary_richness = len(unique_words) / word_count if word_count > 0 else 0

        # Detect potential language patterns
        has_code = any(c in text for c in "{}[]()=;") and (
            "function" in text or "def " in text or "import " in text
        )
        has_json = text.strip().startswith(("{", "["))
        has_xml = text.strip().startswith("<")

        return {
            "char_count": char_count,
            "char_count_no_spaces": char_count_no_spaces,
            "word_count": word_count,
            "sentence_count": sentence_count,
            "avg_word_length": round(avg_word_length, 2),
            "avg_sentence_length": round(avg_sentence_length, 2),
            "unique_words": len(unique_words),
            "vocabulary_richness": round(vocabulary_richness, 3),
            "has_code_structure": has_code,
            "has_json_format": has_json,
            "has_xml_format": has_xml,
            "preview": text[:200] if len(text) > 200 else text,
        }

    async def _extract_image_features(
        self, image_data: Any, format_hint: str | None
    ) -> dict[str, Any]:
        """Extract features from image input."""
        image_bytes = self._decode_image_bytes(image_data)
        mime_type = self._detect_image_mime(image_data)
        base: dict[str, Any] = {
            "format": format_hint or "unknown", "mime_type": mime_type, "size_bytes": len(image_bytes),
        }
        pil_features = self._try_extract_pil(image_bytes)
        llm_description = await self._try_describe_image_llm(image_data)
        return self._merge_image_features(base, pil_features, llm_description)

    def _try_extract_pil(self, image_bytes: bytes) -> dict[str, Any]:
        try:
            from PIL import Image, ImageStat  # noqa: F401
            return self._extract_image_pil(image_bytes)
        except ImportError:
            logger.warning("[%s] PIL/Pillow not available — falling back to metadata",
                           self.agent_id, extra={"event": "perceiver_pil_unavailable"})
        except Exception:
            logger.exception("[%s] PIL image decode failed",
                             self.agent_id, extra={"event": "perceiver_pil_decode_failed"})
        return {}

    async def _try_describe_image_llm(self, image_data: Any) -> str | None:
        if not (self.swarms_agent and self.swarms_agent.llm and isinstance(image_data, str)):
            return None
        try:
            return await asyncio.wait_for(self._describe_image_llm(image_data), timeout=60)
        except TimeoutError:
            logger.warning("[%s] Image LLM analysis timed out",
                           self.agent_id, extra={"event": "perceiver_llm_timeout"})
        except Exception:
            logger.exception("[%s] Image LLM analysis error",
                             self.agent_id, extra={"event": "perceiver_llm_error"})
        return None

    @staticmethod
    def _merge_image_features(
        base: dict[str, Any], pil_features: dict[str, Any], llm_description: str | None
    ) -> dict[str, Any]:
        result = dict(base)
        if pil_features:
            result.update(pil_features)
            result["description"] = llm_description or ""
            result["analyzed_by"] = "pil+llm" if llm_description else "pil"
        elif llm_description:
            result["description"] = llm_description
            result["analyzed_by"] = "llm"
        else:
            result["analyzed_by"] = "metadata"
        return result

    @staticmethod
    def _decode_image_bytes(image_data: Any) -> bytes:
        """Decode ``image_data`` to raw bytes regardless of input format.

        Handles:
        - ``data:image/xxx;base64,...`` data URLs
        - plain base64 strings
        - ``bytes``
        """
        if isinstance(image_data, bytes):
            return image_data
        if not isinstance(image_data, str):
            return b""
        payload = image_data
        if payload.startswith(_DATA_URL_PREFIX):
            # Strip the "data:image/xxx;base64," prefix
            try:
                payload = payload.split(",", 1)[1]
            except IndexError:
                payload = ""
        try:
            return base64.b64decode(payload)
        except Exception:
            return image_data.encode("utf-8")

    @staticmethod
    def _detect_image_mime(image_data: Any) -> str:
        """Infer a MIME type string from the input shape."""
        if isinstance(image_data, str) and image_data.startswith(_DATA_URL_PREFIX):
            try:
                return image_data.split(":")[1].split(";")[0]
            except IndexError:
                return "unknown"
        if isinstance(image_data, bytes):
            # Sniff magic bytes
            if image_data.startswith(b"\xff\xd8\xff"):
                return "image/jpeg"
            if image_data.startswith(b"\x89PNG"):
                return "image/png"
            if image_data.startswith((b"GIF87a", b"GIF89a")):
                return "image/gif"
            return _OCTET_STREAM_MIME
        return "unknown"

    @staticmethod
    def _extract_image_pil(image_bytes: bytes) -> dict[str, Any]:
        """Decode *image_bytes* with Pillow and return structured features.

        Returns a dict with *dimensions*, *mode*, *channels*, *color_stats*,
        and *dominant_color_rgb*.  Raises on decode failure — callers must
        handle PIL unavailability and decode errors.
        """
        import contextlib
        from pathlib import Path

        from PIL import Image, ImageStat

        with tempfile.NamedTemporaryFile(suffix=".img", delete=False) as tmpf:
            tmpf.write(image_bytes)
            tmpf.flush()
            tmp_path = tmpf.name

        try:
            img = Image.open(tmp_path)
            img.load()  # Force decode
            width, height = img.size
            mode = img.mode
            channels = len(img.getbands())

            stat = ImageStat.Stat(img)
            color_stats: dict[str, dict[str, float]] = {}
            band_names = img.getbands()
            for i, band_name in enumerate(band_names):
                mean_val = round(stat.mean[i], 2) if i < len(stat.mean) else 0.0
                stddev_val = round(stat.stddev[i], 2) if i < len(stat.stddev) else 0.0
                color_stats[band_name] = {"mean": mean_val, "stddev": stddev_val}

            # Dominant color (fast quantize to 1 color)
            dominant_color_rgb: list[int] = []
            try:
                quantized = img.quantize(colors=1) if img.mode != "P" else img
                palette = quantized.getpalette()
                if palette:
                    dominant_color_rgb = list(palette[:3])
            except Exception:
                dominant_color_rgb = []

            return {
                "dimensions": {"width": width, "height": height},
                "mode": mode,
                "channels": channels,
                "color_stats": color_stats,
                "dominant_color_rgb": dominant_color_rgb,
            }
        finally:
            with contextlib.suppress(OSError):
                Path(tmp_path).unlink()

    async def _describe_image_llm(self, image_data: str) -> str:
        """Use LLM to describe an image via the provider chain.

        Routes a vision-oriented prompt (including the image data) through
        ``run_with_llm()`` → ModelGarage → provider → model.  Falls back to a
        metadata-only string when the LLM path is unavailable.
        """
        size_bytes = len(image_data.encode())
        fmt = "base64"
        if image_data.startswith(_DATA_URL_PREFIX):
            try:
                mime = image_data.split(":")[1].split(";")[0]
            except IndexError:
                mime = "unknown"
            fmt = mime.split("/")[-1] if "/" in mime else mime

        prompt = (
            "You are a vision analysis assistant.  Describe the following image in detail, "
            "including any text, objects, people, colors, and the overall scene.  "
            f"Image data ({fmt}, {size_bytes} bytes):\n{image_data}"
        )

        try:
            response = await self.run_with_llm(prompt, timeout=60)
            logger.info(
                "[%s] Image description generated via LLM",
                self.agent_id,
                extra={"size_bytes": size_bytes, "format": fmt},
            )
            return response
        except Exception:
            logger.warning(
                "[%s] LLM unavailable for image description",
                self.agent_id,
                extra={
                    "event": "perceiver_llm_unavailable",
                    "size_bytes": size_bytes,
                    "format": fmt,
                },
                exc_info=True,
            )
            return f"[LLM unavailable] Image analysis requested ({size_bytes} bytes, {fmt})"

    @staticmethod
    def _decode_audio_bytes(audio_data: Any) -> tuple[bytes, str]:
        """Decode ``audio_data`` to raw bytes and infer a MIME type.

        Returns ``(bytes, mime_type)``.
        """
        if isinstance(audio_data, bytes):
            return audio_data, _OCTET_STREAM_MIME
        if not isinstance(audio_data, str):
            return b"", "unknown"

        payload = audio_data
        mime_type = _OCTET_STREAM_MIME
        if payload.startswith(_DATA_URL_PREFIX):
            try:
                header, payload = payload.split(",", 1)
                mime_type = header.split(":")[1].split(";")[0]
            except (IndexError, ValueError):
                pass

        try:
            return base64.b64decode(payload), mime_type
        except Exception:
            return audio_data.encode("utf-8"), mime_type

    @staticmethod
    def _audio_suffix_from_format(format_hint: str | None, mime_type: str) -> str:
        """Return a file extension (with dot) for a known audio format or mime."""
        format_lower = (format_hint or "").lower()
        mime_ext_map = {
            "audio/wav": ".wav", "audio/wave": ".wav", "audio/x-wav": ".wav",
            "audio/mpeg": ".mp3", "audio/mp3": ".mp3",
            "audio/ogg": ".ogg", "audio/vorbis": ".ogg",
            "audio/flac": ".flac",
            "audio/aac": ".aac", "audio/x-aac": ".aac",
            "audio/webm": ".webm",
        }
        # Known bare extensions
        if format_lower in {"wav", "mp3", "ogg", "flac", "aac", "webm"}:
            return f".{format_lower}"
        # MIME type lookup (only when we actually have a non-empty mime)
        if mime_type and mime_type != _OCTET_STREAM_MIME:
            for mime_key, ext in mime_ext_map.items():
                if mime_type == mime_key or mime_type.startswith(mime_key):
                    return ext
        if format_lower:
            return f".{format_lower}"
        return ".audio"

    async def _extract_audio_features(
        self, audio_data: Any, format_hint: str | None
    ) -> dict[str, Any]:
        """Extract features from audio input using ffprobe + ffmpeg volumedetect.

        Decodes data URL, bytes, or str; writes to a temp file; runs ffprobe
        for stream metadata and ffmpeg volumedetect for volume stats.
        Falls back to size-only metadata when tools are unavailable or fail.
        """
        import contextlib
        import json
        import re
        from pathlib import Path

        audio_bytes, mime_type = self._decode_audio_bytes(audio_data)
        suffix = self._audio_suffix_from_format(format_hint, mime_type)
        size_bytes = len(audio_bytes)

        base_meta: dict[str, Any] = {
            "format": format_hint or "unknown",
            "mime_type": mime_type,
            "size_bytes": size_bytes,
        }

        tmp_path: str | None = None
        try:
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmpf:
                tmpf.write(audio_bytes)
                tmpf.flush()
                tmp_path = tmpf.name

            # ----- ffprobe stream metadata -----
            proc = await asyncio.create_subprocess_exec(
                "ffprobe", "-v", "quiet", "-print_format", "json",
                "-show_format", "-show_streams", tmp_path,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, stderr = await proc.communicate()
            if proc.returncode != 0:
                raise RuntimeError(
                    f"ffprobe exited {proc.returncode}: {stderr.decode()[:200]}"
                )

            parsed = json.loads(stdout.decode())
            audio_streams = [
                s for s in parsed.get("streams", []) if s.get("codec_type") == "audio"
            ]
            if not audio_streams:
                raise RuntimeError("No audio stream found in file")

            stream = audio_streams[0]
            fmt = parsed.get("format", {})

            # ----- ffmpeg volumedetect -----
            # NOTE: volumedetect filter output goes to stderr at *info* level,
            # so we use "-v info" instead of "-v quiet".
            vproc = await asyncio.create_subprocess_exec(
                "ffmpeg", "-v", "info", "-i", tmp_path,
                "-af", "volumedetect", "-f", "null", "/dev/null",
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            _vout, verr = await vproc.communicate()
            verr_text = verr.decode()

            mean_match = re.search(r"mean_volume:\s*(-?\d+\.?\d*)\s*dB", verr_text)
            max_match = re.search(r"max_volume:\s*(-?\d+\.?\d*)\s*dB", verr_text)

            features: dict[str, Any] = {
                "sample_rate": int(stream.get("sample_rate", 0)),
                "channels": int(stream.get("channels", 0)),
                "codec_name": stream.get("codec_name", "unknown"),
                "bit_rate": int(stream.get("bit_rate", 0)) or int(fmt.get("bit_rate", 0)),
                "duration": float(fmt.get("duration", 0)),
                "mean_volume": float(mean_match.group(1)) if mean_match else None,
                "max_volume": float(max_match.group(1)) if max_match else None,
            }

            result = dict(base_meta)
            result.update(features)
            result["analyzed_by"] = "ffprobe"
            return result

        except Exception as exc:
            logger.warning(
                "[%s] ffprobe audio extraction failed — falling back to metadata",
                self.agent_id,
                extra={"event": "perceiver_ffprobe_failed", "error": str(exc)[:200]},
            )
            return {
                **base_meta,
                "analyzed_by": "metadata",
                "note": f"ffprobe unavailable or failed: {str(exc)[:200]}",
            }

        finally:
            if tmp_path:
                with contextlib.suppress(OSError):
                    Path(tmp_path).unlink()  # noqa: ASYNC240 — cleanup must be sync in finally

    @staticmethod
    def _decode_video_bytes(video_data: Any) -> tuple[bytes, str]:
        """Decode ``video_data`` to raw bytes and infer a MIME type.

        Returns ``(bytes, mime_type)``.
        """
        if isinstance(video_data, bytes):
            return video_data, _OCTET_STREAM_MIME
        if not isinstance(video_data, str):
            return b"", "unknown"

        payload = video_data
        mime_type = _OCTET_STREAM_MIME
        if payload.startswith(_DATA_URL_PREFIX):
            try:
                header, payload = payload.split(",", 1)
                mime_type = header.split(":")[1].split(";")[0]
            except (IndexError, ValueError):
                pass

        try:
            return base64.b64decode(payload), mime_type
        except Exception:
            return video_data.encode("utf-8"), mime_type

    @staticmethod
    def _video_suffix_from_format(format_hint: str | None, mime_type: str) -> str:
        """Return a file extension (with dot) for a known video format or mime."""
        format_lower = (format_hint or "").lower()
        mime_ext_map: dict[str, str] = {
            "video/mp4": ".mp4",
            "video/mpeg": ".mpg",
            "video/avi": ".avi",
            "video/x-msvideo": ".avi",
            "video/quicktime": ".mov",
            "video/webm": ".webm",
            "video/x-matroska": ".mkv",
        }
        if format_lower in {"mp4", "avi", "mov", "webm", "mkv", "mpg", "mpeg"}:
            return f".{format_lower}"
        if mime_type and mime_type != _OCTET_STREAM_MIME:
            for mime_key, ext in mime_ext_map.items():
                if mime_type == mime_key or mime_type.startswith(mime_key):
                    return ext
        if format_lower:
            return f".{format_lower}"
        return ".video"

    @staticmethod
    def _parse_r_frame_rate(r_frame_rate: str | None) -> float | None:
        """Convert ``r_frame_rate`` string (e.g. "30/1" or "30000/1001") to float.

        Returns None when input is empty, malformed, or zero-denominator.
        """
        if not r_frame_rate:
            return None
        try:
            num, denom = r_frame_rate.split("/", 1)
            n = int(num)
            d = int(denom)
            if d == 0:
                return None
            return round(n / d, 2)
        except (ValueError, ZeroDivisionError):
            return None

    async def _extract_video_features(
        self, video_data: Any, format_hint: str | None
    ) -> dict[str, Any]:
        """Extract features from video input using ffprobe.

        Decodes data URL, bytes, or str; writes to a temp file; runs ffprobe
        for stream metadata and format info.  Falls back to size-only metadata
        when ffprobe is unavailable or fails.
        """
        import contextlib
        import json
        from pathlib import Path

        video_bytes, mime_type = self._decode_video_bytes(video_data)
        suffix = self._video_suffix_from_format(format_hint, mime_type)
        size_bytes = len(video_bytes)

        base_meta: dict[str, Any] = {
            "format": format_hint or "unknown",
            "mime_type": mime_type,
            "size_bytes": size_bytes,
        }

        tmp_path: str | None = None
        try:
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmpf:
                tmpf.write(video_bytes)
                tmpf.flush()
                tmp_path = tmpf.name

            # ----- ffprobe stream + format metadata -----
            proc = await asyncio.create_subprocess_exec(
                "ffprobe", "-v", "quiet", "-print_format", "json",
                "-show_format", "-show_streams", tmp_path,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, stderr = await proc.communicate()
            if proc.returncode != 0:
                raise RuntimeError(
                    f"ffprobe exited {proc.returncode}: {stderr.decode()[:200]}"
                )

            parsed = json.loads(stdout.decode())
            video_streams = [
                s for s in parsed.get("streams", []) if s.get("codec_type") == "video"
            ]
            if not video_streams:
                raise RuntimeError("No video stream found in file")

            stream = video_streams[0]
            fmt = parsed.get("format", {})

            r_frame_rate_str: str | None = stream.get("r_frame_rate")
            fps: float | None = self._parse_r_frame_rate(r_frame_rate_str)

            nb_frames_raw: str | None = stream.get("nb_frames")
            nb_frames: int | None
            try:
                nb_frames = int(nb_frames_raw) if nb_frames_raw else None
            except (ValueError, TypeError):
                nb_frames = None

            duration_raw = fmt.get("duration") or stream.get("duration")
            duration: float | None
            try:
                duration = float(duration_raw) if duration_raw else None
            except (ValueError, TypeError):
                duration = None

            bit_rate_raw = stream.get("bit_rate") or fmt.get("bit_rate")
            bit_rate: int | None
            try:
                bit_rate = int(bit_rate_raw) if bit_rate_raw else None
            except (ValueError, TypeError):
                bit_rate = None

            width = stream.get("width")
            height = stream.get("height")

            features: dict[str, Any] = {
                "codec_name": stream.get("codec_name", "unknown"),
                "pix_fmt": stream.get("pix_fmt", "unknown"),
                "bit_rate": bit_rate,
            }
            if width is not None:
                features["width"] = int(width)
            if height is not None:
                features["height"] = int(height)
            if fps is not None:
                features["fps"] = fps
                features["frame_rate_raw"] = r_frame_rate_str
            if nb_frames is not None:
                features["frame_count"] = nb_frames
            if duration is not None:
                features["duration"] = duration

            result = dict(base_meta)
            result.update(features)
            result["analyzed_by"] = "ffprobe"
            return result

        except Exception as exc:
            logger.warning(
                "[%s] ffprobe video extraction failed — falling back to metadata",
                self.agent_id,
                extra={"event": "perceiver_video_ffprobe_failed", "error": str(exc)[:200]},
            )
            return {
                **base_meta,
                "analyzed_by": "metadata",
                "note": f"ffprobe unavailable or failed: {str(exc)[:200]}",
            }

        finally:
            if tmp_path:
                with contextlib.suppress(OSError):
                    Path(tmp_path).unlink()  # noqa: ASYNC240 — sync unlink in async finally (same pattern as T01/T02)

    # --- Text file extensions handled by text-stat parser ---
    _TEXT_FORMATS: ClassVar[set[str]] = {"txt", "md", "json", "xml", "html", "csv", "yaml", "yml"}

    # --- Binary document handlers keyed by extension ---
    _BINARY_PARSERS: ClassVar[dict[str, str]] = {
        "pdf": "PyPDF2",
        "docx": "python-docx",
        "xlsx": "openpyxl",
    }

    @staticmethod
    def _detect_text_structure(text: str, fmt: str) -> dict[str, Any]:
        """Detect document structure from text content.

        Returns a ``structure`` dict with format-specific markers.
        """
        structure: dict[str, Any] = {}

        if fmt == "json":
            # Attempt JSON parse — count top-level keys/items
            import json as _json

            try:
                parsed = _json.loads(text)
                if isinstance(parsed, dict):
                    structure["json_keys"] = len(parsed)
                    structure["json_type"] = "object"
                elif isinstance(parsed, list):
                    structure["json_items"] = len(parsed)
                    structure["json_type"] = "array"
                else:
                    structure["json_type"] = type(parsed).__name__
                structure["json_valid"] = True
            except _json.JSONDecodeError:
                structure["json_valid"] = False

        elif fmt == "xml" or fmt == "html":
            # Count XML/HTML tags with a simple regex
            import re as _re

            tags = _re.findall(r"<\s*/?\s*(\w+)", text)
            if tags:
                tag_counts: dict[str, int] = {}
                for t in tags:
                    tag_counts[t] = tag_counts.get(t, 0) + 1
                structure["tag_count"] = len(tags)
                structure["unique_tags"] = len(tag_counts)
                # Show top-5 most frequent tags
                top_tags = sorted(tag_counts.items(), key=lambda x: -x[1])[:5]
                structure["top_tags"] = dict(top_tags)

        elif fmt == "csv":
            import csv as _csv
            import io as _io

            try:
                reader = _csv.reader(_io.StringIO(text))
                rows = list(reader)
                structure["csv_rows"] = len(rows)
                if rows:
                    structure["csv_columns"] = len(rows[0])
                    structure["csv_headers"] = rows[0]
            except Exception:
                structure["csv_parse_failed"] = True

        elif fmt == "md":
            import re as _re

            # Count markdown headings
            headings = _re.findall(r"^(#{1,6})\s", text, _re.MULTILINE)
            structure["heading_count"] = len(headings)
            # Count fenced code blocks
            fences = _re.findall(r"^```", text, _re.MULTILINE)
            structure["code_block_count"] = len(fences) // 2
            # Count links [text](url)
            links = _re.findall(r"\[([^\]]*)\]\([^)]*\)", text)
            structure["link_count"] = len(links)

        return structure

    async def _extract_document_features(
        self, doc_data: Any, format_hint: str | None
    ) -> dict[str, Any]:
        """Extract features from document input.

        Text formats (txt, md, json, xml, html, csv, yaml):
            word/sentence/line/char counts, structure detection, text preview.

        Binary formats (pdf, docx, xlsx):
            Optional library-based parsing with graceful ImportError fallback.
        """
        fmt = (format_hint or "").lower()

        # ---- Binary format path (bytes or str that could be base64) ----
        if fmt in self._BINARY_PARSERS:
            parser_name = self._BINARY_PARSERS[fmt]

            # Decode bytes if needed
            if isinstance(doc_data, bytes):
                file_bytes = doc_data
            elif isinstance(doc_data, str):
                # Could be a base64-encoded file
                file_bytes = self._decode_image_bytes(doc_data)
            else:
                file_bytes = b""

            return await self._extract_binary_document_features(
                file_bytes, fmt, parser_name
            )

        # ---- Text format path ----
        if isinstance(doc_data, bytes):
            try:
                text = doc_data.decode("utf-8")
            except UnicodeDecodeError:
                text = doc_data.decode("utf-8", errors="replace")
        elif isinstance(doc_data, str):
            text = doc_data
        else:
            text = str(doc_data)

        size_bytes = len(text.encode("utf-8"))
        lines = text.splitlines()
        line_count = len(lines)

        # Word count (non-empty alphanumeric tokens)
        import re as _re

        words = [w for w in _re.split(r"\s+", text) if w]
        word_count = len(words)

        # Sentence count (split on .!? but avoid decimals and abbreviations)
        sentence_tokens = _re.split(r"[.!?]+", text)
        sentences = [s.strip() for s in sentence_tokens if s.strip()]
        sentence_count = len(sentences)

        char_count = len(text)

        # Structure detection
        structure = self._detect_text_structure(text, fmt)

        # Text preview (truncated to 1000 chars)
        text_preview = text[:1000] if len(text) > 1000 else text

        result: dict[str, Any] = {
            "format": fmt or "unknown",
            "size_bytes": size_bytes,
            "word_count": word_count,
            "sentence_count": sentence_count,
            "line_count": line_count,
            "char_count": char_count,
            "text_preview": text_preview,
        }
        if structure:
            result["structure"] = structure

        result["analyzed_by"] = "text-stat"
        return result

    async def _extract_binary_document_features(
        self, file_bytes: bytes, fmt: str, parser_name: str
    ) -> dict[str, Any]:
        """Extract features from binary documents using optional libraries.

        Args:
            file_bytes: Raw file content.
            fmt: Lower-cased format extension (pdf, docx, xlsx).
            parser_name: Canonical library name for ImportError messages.

        Returns:
            Dict with format-specific fields plus ``analyzed_by``.
        """
        import contextlib
        import tempfile
        from pathlib import Path

        base: dict[str, Any] = {
            "format": fmt,
            "size_bytes": len(file_bytes),
        }

        tmp_path: str | None = None
        try:
            with tempfile.NamedTemporaryFile(
                suffix=f".{fmt}", delete=False
            ) as tmpf:
                tmpf.write(file_bytes)
                tmpf.flush()
                tmp_path = tmpf.name

            if fmt == "pdf":
                return await self._extract_pdf_features(tmp_path, base)
            if fmt == "docx":
                return await self._extract_docx_features(tmp_path, base)
            if fmt == "xlsx":
                return await self._extract_xlsx_features(tmp_path, base)

            # Unknown binary format — shouldn't reach here
            return {**base, "analyzed_by": "metadata", "note": f"Unknown binary format: {fmt}"}

        except ImportError as exc:
            logger.warning(
                "[%s] %s not available for %s parsing — falling back to metadata",
                self.agent_id,
                parser_name,
                fmt,
                extra={
                    "event": "perceiver_doc_lib_unavailable",
                    "parser": parser_name,
                    "format": fmt,
                },
            )
            return {
                **base,
                "analyzed_by": "metadata",
                "note": f"{parser_name} not installed: {exc}",
            }
        except Exception as exc:
            logger.warning(
                "[%s] %s parse failed for %s — falling back to metadata",
                self.agent_id,
                parser_name,
                fmt,
                extra={
                    "event": "perceiver_doc_parse_failed",
                    "parser": parser_name,
                    "format": fmt,
                    "error": str(exc)[:200],
                },
            )
            return {
                **base,
                "analyzed_by": "metadata",
                "note": f"{parser_name} parse error: {str(exc)[:200]}",
            }
        finally:
            if tmp_path:
                with contextlib.suppress(OSError):
                    Path(tmp_path).unlink()  # noqa: ASYNC240

    async def _extract_pdf_features(
        self, path: str, base: dict[str, Any]
    ) -> dict[str, Any]:
        """Extract PDF features using PyPDF2."""
        from PyPDF2 import PdfReader  # type: ignore[import-untyped]

        reader = PdfReader(path)
        page_count = len(reader.pages)

        # Extract text from first 5 pages (limit for performance)
        pages_text: list[str] = []
        for _i, page in enumerate(reader.pages[:5]):
            try:
                extracted = page.extract_text()
                if extracted:
                    pages_text.append(extracted.strip())
            except Exception as exc:
                logger.debug(
                    "[%s] PyPDF2 page text extraction skipped",
                    "perceiver",
                    extra={"error": str(exc)[:200]},
                )

        preview = "\n".join(pages_text)[:1000] if pages_text else ""

        result: dict[str, Any] = dict(base)
        result.update({
            "page_count": page_count,
            "analyzed_by": "PyPDF2",
        })
        if preview:
            result["text_preview"] = preview

        # Metadata (title, author, etc.)
        meta = reader.metadata
        if meta:
            doc_meta: dict[str, Any] = {}
            for key in ("title", "author", "creator", "producer", "subject"):
                val = getattr(meta, key, None)
                if val:
                    doc_meta[key] = str(val)
            if doc_meta:
                result["metadata"] = doc_meta

        return result

    async def _extract_docx_features(
        self, path: str, base: dict[str, Any]
    ) -> dict[str, Any]:
        """Extract DOCX features using python-docx."""
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        paragraph_count = len(paragraphs)

        # Concatenate first paragraphs for preview (up to 1000 chars)
        preview_parts: list[str] = []
        char_budget = 1000
        for p_text in paragraphs:
            preview_parts.append(p_text)
            char_budget -= len(p_text) + 1
            if char_budget <= 0:
                break
        preview = "\n".join(preview_parts)[:1000]

        result: dict[str, Any] = dict(base)
        result.update({
            "paragraph_count": paragraph_count,
            "analyzed_by": "python-docx",
        })
        if preview:
            result["text_preview"] = preview

        # Tables count
        tables = doc.tables
        if tables:
            result["table_count"] = len(tables)

        # Optional core properties
        try:
            props = doc.core_properties
            doc_props: dict[str, Any] = {}
            for attr in ("title", "author", "subject", "keywords"):
                val = getattr(props, attr, None)
                if val:
                    doc_props[attr] = str(val)
            if doc_props:
                result["metadata"] = doc_props
        except Exception as exc:
            logger.debug(
                "[%s] DOCX core properties unavailable",
                self.agent_id,
                extra={"error": str(exc)[:200]},
            )

        return result

    async def _extract_xlsx_features(
        self, path: str, base: dict[str, Any]
    ) -> dict[str, Any]:
        """Extract XLSX features using openpyxl (read-only mode)."""
        from openpyxl import load_workbook  # type: ignore[import-untyped]

        wb = load_workbook(path, read_only=True, data_only=True)
        sheet_names = wb.sheetnames
        sheets_info: list[dict[str, Any]] = []

        for sname in sheet_names:
            ws = wb[sname]
            info: dict[str, Any] = {"name": sname}
            # Dimensions / max_row / max_column may not be available in read_only mode
            try:
                dimensions = getattr(ws, "dimensions", None)
                if dimensions:
                    info["dimensions"] = str(dimensions)
            except Exception as exc:
                logger.debug(
                    "[%s] XLSX dimensions unavailable for sheet %s: %s",
                    self.agent_id,
                    sname,
                    str(exc)[:200],
                )
            try:
                if ws.max_row:
                    info["rows"] = ws.max_row
            except Exception as exc:
                logger.debug(
                    "[%s] XLSX max_row unavailable for sheet %s: %s",
                    self.agent_id,
                    sname,
                    str(exc)[:200],
                )
            try:
                if ws.max_column:
                    info["columns"] = ws.max_column
            except Exception as exc:
                logger.debug(
                    "[%s] XLSX max_column unavailable for sheet %s: %s",
                    self.agent_id,
                    sname,
                    str(exc)[:200],
                )
            sheets_info.append(info)

        wb.close()

        result: dict[str, Any] = dict(base)
        result.update({
            "sheet_names": sheet_names,
            "sheet_count": len(sheet_names),
            "sheets": sheets_info,
            "analyzed_by": "openpyxl",
        })

        return result

    def _extract_sensor_features(self, sensor_data: dict[str, Any]) -> dict[str, Any]:
        """Extract features from sensor data."""
        if not isinstance(sensor_data, dict):
            return {"error": "Sensor data must be a dictionary"}

        # Extract basic statistics from numeric values
        numeric_values = []
        for value in sensor_data.values():
            if isinstance(value, (int, float)):
                numeric_values.append(value)  # noqa: PERF401

        stats = {}
        if numeric_values:
            stats = {
                "min": min(numeric_values),
                "max": max(numeric_values),
                "avg": sum(numeric_values) / len(numeric_values),
                "count": len(numeric_values),
            }

        return {
            "keys": list(sensor_data.keys()),
            "numeric_stats": stats,
            "analyzed_by": "sensor_parser",
        }

    def _assess_input_quality(
        self, input_data: Any, modality: str, features: dict[str, Any]  # noqa: ARG002
    ) -> float:
        """Assess input quality (0-1 score)."""
        quality_score = 1.0

        # Check for errors in features
        if "error" in features:
            quality_score -= 0.5

        # Modality-specific quality checks
        if modality == ModalityType.TEXT.value:
            # Check for meaningful content
            word_count = features.get("word_count", 0)
            if word_count < 3:
                quality_score -= 0.3
            if word_count > 100000:
                quality_score -= 0.2  # Very long text may need chunking

        elif modality == ModalityType.IMAGE.value:
            size_bytes = features.get("size_bytes", 0)
            if size_bytes < 1000:
                quality_score -= 0.4  # Very small image
            if size_bytes > 50 * 1024 * 1024:
                quality_score -= 0.3  # Very large image

        # Ensure score stays in valid range
        return max(0.0, min(1.0, quality_score))

    def _cache_features(
        self, input_id: str, modality: str, features: dict[str, Any], metadata: dict[str, Any]
    ) -> None:
        """Cache features for cross-modal correlation."""
        if not self.enable_cross_modal:
            return

        self.feature_cache[input_id] = {
            "modality": modality,
            "features": features,
            "metadata": metadata,
            "timestamp": datetime.now(UTC).isoformat(),
        }

        # Enforce cache size limit
        if len(self.feature_cache) > self.feature_cache_size:
            # Remove oldest entries
            sorted_keys = sorted(
                self.feature_cache.keys(), key=lambda k: self.feature_cache[k]["timestamp"]
            )
            for key in sorted_keys[:100]:  # Remove 100 oldest
                del self.feature_cache[key]

    async def _store_in_historian(
        self, input_id: str, modality: str, features: dict[str, Any], metadata: dict[str, Any]
    ) -> None:
        """Store processed input in Historian memory."""
        try:
            await self.send(
                topic="actor:historian",
                content={
                    "message_type": "store_memory",
                    "content": {
                        "input_id": input_id,
                        "modality": modality,
                        "features": features,
                        "metadata": metadata,
                    },
                    "metadata": {
                        "source": "perceiver",
                        "timestamp": datetime.now(UTC).isoformat(),
                    },
                },
            )
        except Exception:
            logger.warning(f"[{self.agent_id}] Failed to store in Historian: {e}")

    async def _handle_extract_features(self, message: ActorMessage) -> None:
        """
        Extract features from existing input.

        Args:
            message: ActorMessage with content containing:
                - input_id: ID of previously processed input
                - modality: Optional modality override
        """
        try:
            content = message.content
            input_id = content.get("input_id")

            if not input_id:
                await self._send_error_response(message, "Missing input_id")
                return

            cached = self.feature_cache.get(input_id)
            if not cached:
                await self._send_error_response(message, f"Input not found: {input_id}")
                return

            await self.send(
                topic=content.get("reply_to"),
                content={
                    "message_type": "features_result",
                    "input_id": input_id,
                    **cached,
                },
                correlation_id=message.correlation_id,
            )

        except Exception as e:
            logger.exception(f"[{self.agent_id}] Feature extraction failed: {e}")
            await self._send_error_response(message, f"Feature extraction failed: {e}")

    async def _handle_classify_modality(self, message: ActorMessage) -> None:
        """
        Classify input modality.

        Args:
            message: ActorMessage with content containing:
                - input_data: Data to classify
                - format_hint: Optional format hint
        """
        try:
            content = message.content
            input_data = content.get("input_data")
            format_hint = content.get("format_hint")

            if input_data is None:
                await self._send_error_response(message, "Missing input_data")
                return

            modality = self._detect_modality(input_data, format_hint)

            await self.send(
                topic=content.get("reply_to"),
                content={
                    "message_type": "modality_result",
                    "modality": modality,
                    "confidence": 0.8,  # Placeholder confidence
                },
                correlation_id=message.correlation_id,
            )

        except Exception as e:
            logger.exception(f"[{self.agent_id}] Modality classification failed: {e}")
            await self._send_error_response(message, f"Classification failed: {e}")

    async def _handle_assess_quality(self, message: ActorMessage) -> None:
        """
        Assess input quality.

        Args:
            message: ActorMessage with content containing:
                - input_id: ID of input to assess
        """
        try:
            content = message.content
            input_id = content.get("input_id")

            if not input_id:
                await self._send_error_response(message, "Missing input_id")
                return

            cached = self.feature_cache.get(input_id)
            if not cached:
                await self._send_error_response(message, f"Input not found: {input_id}")
                return

            quality_score = self._assess_input_quality(None, cached["modality"], cached["features"])

            await self.send(
                topic=content.get("reply_to"),
                content={
                    "message_type": "quality_result",
                    "input_id": input_id,
                    "quality_score": quality_score,
                },
                correlation_id=message.correlation_id,
            )

        except Exception as e:
            logger.exception(f"[{self.agent_id}] Quality assessment failed: {e}")
            await self._send_error_response(message, f"Quality assessment failed: {e}")

    async def _handle_get_processing_stats(self, message: ActorMessage) -> None:
        """
        Get processing statistics.

        Args:
            message: ActorMessage
        """
        try:
            stats = {
                "inputs_processed": self.inputs_processed.copy(),
                "total_features_extracted": self.total_features_extracted,
                "quality_rejections": self.quality_rejections,
                "cache_size": len(self.feature_cache),
                "cross_modal_correlations": len(self.cross_modal_correlations),
                "timestamp": datetime.now(UTC).isoformat(),
            }

            await self.send(
                topic=message.content.get("reply_to"),
                content={
                    "message_type": "stats_result",
                    **stats,
                },
                correlation_id=message.correlation_id,
            )

        except Exception as e:
            logger.exception(f"[{self.agent_id}] Stats retrieval failed: {e}")
            await self._send_error_response(message, f"Stats retrieval failed: {e}")

    async def _handle_correlate_modalities(self, message: ActorMessage) -> None:
        """
        Find cross-modal correlations.

        Args:
            message: ActorMessage with content containing:
                - input_ids: List of input IDs to correlate
        """
        try:
            content = message.content
            input_ids = content.get("input_ids", [])

            if len(input_ids) < 2:
                await self._send_error_response(
                    message, "At least 2 input_ids required for correlation"
                )
                return

            correlations = []
            for i, id1 in enumerate(input_ids):
                for id2 in input_ids[i + 1 :]:
                    data1 = self.feature_cache.get(id1)
                    data2 = self.feature_cache.get(id2)

                    if data1 and data2:
                        correlation = {
                            "input_1": id1,
                            "input_2": id2,
                            "modality_1": data1["modality"],
                            "modality_2": data2["modality"],
                            "timestamp_1": data1["timestamp"],
                            "timestamp_2": data2["timestamp"],
                        }
                        correlations.append(correlation)

                        # Store for future reference
                        self.cross_modal_correlations.append(correlation)

            await self.send(
                topic=content.get("reply_to"),
                content={
                    "message_type": "correlation_result",
                    "correlations": correlations,
                    "count": len(correlations),
                },
                correlation_id=message.correlation_id,
            )

        except Exception as e:
            logger.exception(f"[{self.agent_id}] Correlation analysis failed: {e}")
            await self._send_error_response(message, f"Correlation failed: {e}")

    # Session 44: Collective Learning, Consensus Deliberation, and Memory Optimization
    # integration methods now provided by DeliberationMixin, LearningMixin,
    # MemoryMixin, and PatternMixin.

    async def _send_error_response(self, message: ActorMessage, error: str) -> None:
        """Send error response."""
        if message.content.get("reply_to"):
            await self.send(
                topic=message.content["reply_to"],
                content={
                    "message_type": "error_response",
                    "error": error,
                    "original_message_type": message.message_type,
                },
                correlation_id=message.correlation_id,
            )
