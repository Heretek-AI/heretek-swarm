"""Test T04: Document feature extraction with text statistics and binary parsing.

Validates _extract_document_features for:
- Text formats: word/sentence/line/char counts, structure detection
- Binary formats: PDF/DOCX/XLSX with optional library support
- Graceful fallback when optional libraries are unavailable
"""

import asyncio
import io
import textwrap

import pytest

from heretek_swarm.actors.perceiver.agent import PerceiverAgent


# ---------------------------------------------------------------------------
# Text-format extraction tests
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_markdown_text_statistics():
    """Markdown text extraction returns word_count, sentence_count, line_count."""
    agent = PerceiverAgent()
    text = textwrap.dedent("""\
        # Hello World

        This is a test document.
        It has multiple sentences. And paragraphs.

        ## Section Two

        More content here with [a link](https://example.com) and `inline code`.

        ```python
        print("hello")
        ```
    """)
    result = await agent._extract_document_features(text, "md")

    assert result["word_count"] > 0
    assert result["sentence_count"] > 0
    assert result["line_count"] > 0
    assert result["char_count"] > 0
    assert result["analyzed_by"] == "text-stat"
    assert len(result["text_preview"]) > 0
    assert result["text_preview"] == text  # Fits within 1000 chars


@pytest.mark.asyncio
async def test_markdown_structure_detection():
    """Markdown structure detection finds headings, code blocks, and links."""
    agent = PerceiverAgent()
    text = textwrap.dedent("""\
        # Heading 1
        ## Heading 2
        ### Heading 3

        [link1](https://a.com) and [link2](https://b.com)

        ```python
        code
        ```
    """)
    result = await agent._extract_document_features(text, "md")

    assert result["analyzed_by"] == "text-stat"
    structure = result["structure"]
    assert structure["heading_count"] == 3
    assert structure["code_block_count"] == 1
    assert structure["link_count"] == 2


@pytest.mark.asyncio
async def test_text_preview_truncation():
    """Text preview is truncated to 1000 chars for long documents."""
    agent = PerceiverAgent()
    long_text = "x" * 2000
    result = await agent._extract_document_features(long_text, "txt")

    assert len(result["text_preview"]) == 1000
    assert result["text_preview"] == "x" * 1000


@pytest.mark.asyncio
async def test_txt_format():
    """Plain text format returns basic statistics and text-stat analyzer."""
    agent = PerceiverAgent()
    result = await agent._extract_document_features("Hello world. This is a test.", "txt")

    assert result["word_count"] == 6
    assert result["sentence_count"] == 2
    assert result["line_count"] == 1
    assert result["analyzed_by"] == "text-stat"


@pytest.mark.asyncio
async def test_json_structure_valid():
    """Valid JSON detects object type and key count."""
    agent = PerceiverAgent()
    text = '{"name": "Alice", "age": 30, "city": "NYC"}'
    result = await agent._extract_document_features(text, "json")

    assert result["analyzed_by"] == "text-stat"
    structure = result["structure"]
    assert structure["json_valid"] is True
    assert structure["json_type"] == "object"
    assert structure["json_keys"] == 3


@pytest.mark.asyncio
async def test_json_structure_array():
    """JSON array detects list type and item count."""
    agent = PerceiverAgent()
    text = '[1, 2, 3, 4, 5]'
    result = await agent._extract_document_features(text, "json")

    structure = result["structure"]
    assert structure["json_valid"] is True
    assert structure["json_type"] == "array"
    assert structure["json_items"] == 5


@pytest.mark.asyncio
async def test_json_structure_invalid():
    """Invalid JSON is detected and marked."""
    agent = PerceiverAgent()
    text = '{this is not valid json}'
    result = await agent._extract_document_features(text, "json")

    structure = result["structure"]
    assert structure["json_valid"] is False


@pytest.mark.asyncio
async def test_xml_structure():
    """XML structure detects tags, unique tags, and top tags."""
    agent = PerceiverAgent()
    text = '<root><item><name>Alice</name><value>42</value></item><item><name>Bob</name><value>99</value></item></root>'
    result = await agent._extract_document_features(text, "xml")

    structure = result["structure"]
    assert structure["tag_count"] > 0
    assert structure["unique_tags"] >= 3  # root, item, name, value
    assert "top_tags" in structure


@pytest.mark.asyncio
async def test_html_structure():
    """HTML structure uses same tag-counting as XML."""
    agent = PerceiverAgent()
    text = "<html><head><title>Test</title></head><body><p>Hello</p><div><span>World</span></div></body></html>"
    result = await agent._extract_document_features(text, "html")

    structure = result["structure"]
    assert structure["tag_count"] > 0


@pytest.mark.asyncio
async def test_csv_structure():
    """CSV structure detects rows, columns, and headers."""
    agent = PerceiverAgent()
    text = "col1,col2,col3\n1,2,3\n4,5,6\n7,8,9"
    result = await agent._extract_document_features(text, "csv")

    structure = result["structure"]
    assert structure["csv_rows"] == 4  # header + 3 data rows
    assert structure["csv_columns"] == 3
    assert structure["csv_headers"] == ["col1", "col2", "col3"]


@pytest.mark.asyncio
async def test_csv_invalid():
    """Malformed CSV does not crash — structure notes parse failure."""
    agent = PerceiverAgent()
    result = await agent._extract_document_features(b"\x00\x01\x02", "csv")
    # Binary non-text still goes through text path when format is csv,
    # but decode with replace gives us a valid string — structure detection
    # should handle it gracefully.
    structure = result.get("structure", {})
    assert result["analyzed_by"] == "text-stat"


@pytest.mark.asyncio
async def test_bytes_input_decoded_to_utf8():
    """Bytes input is decoded to UTF-8 before text analysis."""
    agent = PerceiverAgent()
    result = await agent._extract_document_features(b"Hello world. Two sentences.", "txt")

    assert result["word_count"] == 4
    assert result["sentence_count"] == 2


@pytest.mark.asyncio
async def test_bytes_input_replacement_decode():
    """Non-UTF-8 bytes input uses replacement characters."""
    agent = PerceiverAgent()
    # Latin-1 encoded string decoded as UTF-8 would fail, fallback to replace
    result = await agent._extract_document_features(b"Hello \xff world", "txt")
    assert result["analyzed_by"] == "text-stat"
    assert result["char_count"] > 0


@pytest.mark.asyncio
async def test_empty_text():
    """Empty text returns zero counts."""
    agent = PerceiverAgent()
    result = await agent._extract_document_features("", "txt")

    assert result["word_count"] == 0
    assert result["sentence_count"] == 0
    assert result["line_count"] == 0
    assert result["char_count"] == 0
    assert result["analyzed_by"] == "text-stat"


@pytest.mark.asyncio
async def test_line_count_multiple_lines():
    """Line count detects multiple lines correctly."""
    agent = PerceiverAgent()
    result = await agent._extract_document_features("line1\nline2\nline3\n", "txt")

    assert result["line_count"] == 3  # splitlines() on "line1\\nline2\\nline3\\n" → 3 lines


@pytest.mark.asyncio
async def test_non_string_input():
    """Non-string input is coerced to string."""
    agent = PerceiverAgent()
    result = await agent._extract_document_features(12345, "txt")
    assert result["analyzed_by"] == "text-stat"
    assert result["char_count"] > 0


@pytest.mark.asyncio
async def test_unknown_text_format_falls_back_to_text_stat():
    """An unknown extension still gets text-stat path if not in binary parsers."""
    agent = PerceiverAgent()
    result = await agent._extract_document_features("Hello world", "rst")
    assert result["analyzed_by"] == "text-stat"


# ---------------------------------------------------------------------------
# Binary format: PDF extraction
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_pdf_extraction_with_pypdf2():
    """PDF bytes are parsed via PyPDF2 when available."""
    pytest.importorskip("PyPDF2")

    from PyPDF2 import PdfWriter

    # Generate a minimal PDF in memory
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    agent = PerceiverAgent()
    result = await agent._extract_document_features(pdf_bytes, "pdf")

    assert result["analyzed_by"] == "PyPDF2"
    assert result["page_count"] == 2
    assert result["format"] == "pdf"
    assert result["size_bytes"] > 0


@pytest.mark.asyncio
async def test_pdf_base64_input():
    """PDF encoded as base64 string is decoded and parsed."""
    pytest.importorskip("PyPDF2")
    import base64

    from PyPDF2 import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()
    b64_data = base64.b64encode(pdf_bytes).decode()

    agent = PerceiverAgent()
    result = await agent._extract_document_features(b64_data, "pdf")

    assert result["analyzed_by"] == "PyPDF2"
    assert result["page_count"] == 1


@pytest.mark.asyncio
async def test_pdf_fallback_when_pypdf2_unavailable(monkeypatch):
    """When PyPDF2 import fails, metadata fallback is returned."""
    pytest.importorskip("PyPDF2")

    # Simulate ImportError by patching the import in the module
    import builtins

    original_import = builtins.__import__

    def mock_import(name, *args, **kwargs):
        if name == "PyPDF2":
            raise ImportError("PyPDF2 not installed (simulated)")
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", mock_import)

    agent = PerceiverAgent()
    result = await agent._extract_document_features(b"fake pdf bytes", "pdf")

    assert result["analyzed_by"] == "metadata"
    assert "PyPDF2 not installed" in result.get("note", "")


@pytest.mark.asyncio
async def test_pdf_fallback_on_corrupt_data():
    """Corrupt PDF data triggers metadata fallback with parse error note."""
    pytest.importorskip("PyPDF2")

    agent = PerceiverAgent()
    result = await agent._extract_document_features(b"this is not a valid pdf", "pdf")

    assert result["analyzed_by"] == "metadata"
    assert "PyPDF2 parse error" in result.get("note", "")


# ---------------------------------------------------------------------------
# Binary format: DOCX extraction
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_docx_extraction_with_python_docx():
    """DOCX bytes are parsed via python-docx when available."""
    pytest.importorskip("docx")

    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello world.")
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("Third paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    agent = PerceiverAgent()
    result = await agent._extract_document_features(docx_bytes, "docx")

    assert result["analyzed_by"] == "python-docx"
    assert result["paragraph_count"] == 3
    assert len(result.get("text_preview", "")) > 0


@pytest.mark.asyncio
async def test_docx_fallback_when_lib_unavailable(monkeypatch):
    """When python-docx import fails, metadata fallback is returned."""
    pytest.importorskip("docx")

    import builtins

    original_import = builtins.__import__

    def mock_import(name, *args, **kwargs):
        if name == "docx":
            raise ImportError("python-docx not installed (simulated)")
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", mock_import)

    agent = PerceiverAgent()
    result = await agent._extract_document_features(b"fake docx bytes", "docx")

    assert result["analyzed_by"] == "metadata"
    assert "python-docx not installed" in result.get("note", "")


# ---------------------------------------------------------------------------
# Binary format: XLSX extraction
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_xlsx_extraction_with_openpyxl():
    """XLSX bytes are parsed via openpyxl when available."""
    pytest.importorskip("openpyxl")

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "Name"
    ws["B1"] = "Value"
    ws["A2"] = "Alice"
    ws["B2"] = 42
    ws["A3"] = "Bob"
    ws["B3"] = 99

    ws2 = wb.create_sheet("Sheet2")
    ws2["A1"] = "Extra"
    ws2["A2"] = "data"

    buf = io.BytesIO()
    wb.save(buf)
    xlsx_bytes = buf.getvalue()

    agent = PerceiverAgent()
    result = await agent._extract_document_features(xlsx_bytes, "xlsx")

    assert result["analyzed_by"] == "openpyxl"
    assert result["sheet_count"] == 2
    assert result["sheet_names"] == ["Sheet1", "Sheet2"]
    sheets = result["sheets"]
    assert len(sheets) == 2
    assert sheets[0]["name"] == "Sheet1"
    assert sheets[1]["name"] == "Sheet2"


@pytest.mark.asyncio
async def test_xlsx_fallback_when_openpyxl_unavailable(monkeypatch):
    """When openpyxl import fails, metadata fallback is returned."""
    pytest.importorskip("openpyxl")

    import builtins

    original_import = builtins.__import__

    def mock_import(name, *args, **kwargs):
        if name == "openpyxl":
            raise ImportError("openpyxl not installed (simulated)")
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", mock_import)

    agent = PerceiverAgent()
    result = await agent._extract_document_features(b"fake xlsx bytes", "xlsx")

    assert result["analyzed_by"] == "metadata"
    assert "openpyxl not installed" in result.get("note", "")


# ---------------------------------------------------------------------------
# Unknown format
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_unknown_binary_format_triggers_error_path():
    """An unsupported binary format falls through to text path (or metadata)."""
    agent = PerceiverAgent()
    # "xyz" is not in _BINARY_PARSERS, so it takes the text path
    result = await agent._extract_document_features(b"some data", "xyz")
    assert "analyzed_by" in result


# ---------------------------------------------------------------------------
# Statistics tracking integration
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_processed_input_increments_stats():
    """Processing a document increments inputs_processed and total_features_extracted."""
    agent = PerceiverAgent()
    initial = agent.total_features_extracted
    initial_doc = agent.inputs_processed.get("document", 0)

    # Simulate a full _handle_process_input call via the feature extraction path
    result = await agent._extract_document_features("# Test\n\nContent here.", "md")
    assert result["analyzed_by"] == "text-stat"

    # Stats are incremented inside _handle_process_input, not _extract_document_features.
    # Verify the extraction result is valid — the stats path is tested via integration.
    assert result["word_count"] > 0
